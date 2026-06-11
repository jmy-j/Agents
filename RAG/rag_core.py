# -*- coding: utf-8 -*-
"""
RAG 核心引擎 — 文档解析、分块、向量化存储、检索、LLM 生成回答
"""
import os
from langchain_ollama import OllamaEmbeddings, ChatOllama
from langchain_chroma import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
import docx

# ============================================================
# 配置 — 可按需修改
# ============================================================
EMBEDDING_MODEL = "nomic-embed-text"   # embedding 模型，把文字变成向量
CHAT_MODEL = "qwen:0.5b"               # 对话模型，生成回答
CHUNK_SIZE = 500                       # 每个文本块的字数上限
CHUNK_OVERLAP = 50                     # 相邻块之间重叠的字数（防止语义被切断）
SEARCH_K = 4                           # 每次检索返回最相关的块数
CHROMA_DIR = "./chroma_db"             # 向量库本地存储路径


class RAGEngine:
    """封装了文档解析→分块→向量化→检索→生成的全流程"""

    def __init__(self):
        # ---- 初始化三个核心组件 ----
        self.embeddings = OllamaEmbeddings(model=EMBEDDING_MODEL)
        self.llm = ChatOllama(model=CHAT_MODEL, temperature=0.3)

        # 文本分块器：按段落→句子→空格→字符的优先级切割
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            separators=["\n\n", "\n", "。", ".", "！", "!", "？", "?", "；", ";", " ", ""],
        )

        # 向量库：如果本地已有数据就加载，否则等添加文档时再创建
        self.vectorstore = None
        self._load_existing_db()

    # ============================================================
    # 向量库
    # ============================================================

    def _load_existing_db(self):
        """尝试加载本地已有的向量数据库"""
        if os.path.exists(CHROMA_DIR) and os.listdir(CHROMA_DIR):
            try:
                self.vectorstore = Chroma(
                    persist_directory=CHROMA_DIR,
                    embedding_function=self.embeddings,
                )
                return True
            except Exception:
                pass
        return False

    # ============================================================
    # 文档解析 + 入库
    # ============================================================

    def parse_docx(self, file_path: str) -> list[Document]:
        """解析一个 .docx 文件，返回分块后的 Document 列表"""
        doc = docx.Document(file_path)

        # 提取所有段落文字
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        raw_text = "\n".join(paragraphs)

        if not raw_text:
            return []

        # 分块
        chunks = self.text_splitter.split_text(raw_text)
        filename = os.path.basename(file_path)

        documents = [
            Document(
                page_content=chunk,
                metadata={"source": filename, "chunk_index": i},
            )
            for i, chunk in enumerate(chunks)
        ]
        return documents

    def add_document(self, file_path: str) -> int:
        """将一个 .docx 文件解析并存入向量库，返回存入的文本块数"""
        documents = self.parse_docx(file_path)
        if not documents:
            return 0

        if self.vectorstore is None:
            # 首次添加：创建新库
            self.vectorstore = Chroma.from_documents(
                documents=documents,
                embedding=self.embeddings,
                persist_directory=CHROMA_DIR,
            )
        else:
            # 已有库：追加
            self.vectorstore.add_documents(documents)

        return len(documents)

    # ============================================================
    # 检索 + 问答
    # ============================================================

    def search(self, query: str) -> list[Document]:
        """用问题去向量库检索最相关的文本块"""
        if self.vectorstore is None:
            return []
        return self.vectorstore.similarity_search(query, k=SEARCH_K)

    def ask(self, question: str) -> str:
        """完整的 RAG 问答：检索 → 拼接 prompt → LLM 生成"""
        docs = self.search(question)
        if not docs:
            return "向量库中没有找到相关文档，请先在侧边栏上传 .docx 文件。"

        # 拼接检索到的内容作为"参考资料"
        context_parts = []
        for doc in docs:
            source = doc.metadata.get("source", "未知文档")
            context_parts.append(f"[来源: {source}]\n{doc.page_content}")
        context = "\n\n---\n\n".join(context_parts)

        # 构建 prompt：要求 LLM 严格依据文档回答
        prompt = ChatPromptTemplate.from_template(
            "你是一个文档助手，请根据以下文档内容回答用户的问题。\n"
            "如果文档内容不足以回答问题，请如实说明「文档中没有相关信息」，不要编造。\n"
            "\n"
            "## 文档内容：\n"
            "{context}\n"
            "\n"
            "## 用户问题：\n"
            "{question}\n"
            "\n"
            "## 回答："
        )

        chain = prompt | self.llm | StrOutputParser()
        return chain.invoke({"context": context, "question": question})
