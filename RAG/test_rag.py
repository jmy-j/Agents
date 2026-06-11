"""端到端测试 RAG 流程"""
import tempfile, os, sys
sys.path.insert(0, "D:\\python-project\\py-project-1\\RAG")

from rag_core import RAGEngine
from docx import Document

# 创建测试 .docx
doc = Document()
doc.add_paragraph("Python is a popular programming language created by Guido van Rossum in 1991.")
doc.add_paragraph("RAG stands for Retrieval-Augmented Generation, a technique that helps LLMs answer questions based on provided documents.")
tmp_path = os.path.join(tempfile.gettempdir(), "test_rag_demo.docx")
doc.save(tmp_path)
print("测试文件已创建:", tmp_path)

# 测试 RAG 引擎
rag = RAGEngine()
count = rag.add_document(tmp_path)
print("入库块数:", count)

# 测试检索
docs = rag.search("Who created Python?")
print("检索到的块数:", len(docs))
if docs:
    print("第一块内容:", docs[0].page_content[:100])

# 测试问答
answer = rag.ask("Who created Python?")
print("回答:", answer[:300] if answer else "(空)")

os.unlink(tmp_path)
print("\n测试完成!")
